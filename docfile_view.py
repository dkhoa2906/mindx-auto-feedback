from docx import Document
import os
#

doc = Document(os.path.join("document_templates", "coding.docx"))

for i, para in enumerate(doc.paragraphs):
    print(f"\n--- Paragraph {i+1} ---")
    for j, run in enumerate(para.runs):
        print(f"Run {j+1}: '{run.text}'")

for table_index, table in enumerate(doc.tables):
    print(f"\n--- Table {table_index+1} ---")
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            print(f"\n--- Table Row {row_index+1}, Cell {cell_index+1} ---")
            # Loop through paragraphs in each cell
            for para_index, para in enumerate(cell.paragraphs):
                print(f"\n--- Paragraph {para_index+1} in Cell {cell_index+1} ---")
                for run_index, run in enumerate(para.runs):
                    print(f"Run {run_index+1}: '{run.text}'")