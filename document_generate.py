from docx import Document

import os

from docx.shared import Pt

from document_manage import create_output_path, create_from_template


def fill_in_form(student_name, student_age, subject, trial_date, trial_mentor, trial_place, points, feedback):
    document = Document(create_from_template(student_name, subject, trial_date))

    replacements = {
        "{{student_name}}": student_name,
        "{{student_age}}": student_age,
        "{{subject}}": subject,
        "{{trial_date}}": trial_date.strftime("%d/%m/%Y"),
        "{{trial_mentor}}": trial_mentor,
        "{{trial_place}}": trial_place,
    }

    # Add points to replacements
    for i, point in enumerate(points):
        row = i + 1
        for col in range(1, 6):
            placeholder = f"{{{{p{row}{col}}}}}"
            replacements[placeholder] = "✅" if col == point else ""

    # Handle {{p95}} error
    replacements["{{p9}}"] = replacements["{{p95}}"]

    # Calculate average point
    avg_point = (sum(points) / 9)
    replacements["{{avg_point}}"] = f"{avg_point:.1f}"

    # Add types to replacements
    if 4 <= avg_point <= 5:
        replacements["{{type_1}}"] = "✅"
        replacements["{{type_2}}"] = replacements["{{type_3}}"] = replacements["{{type_4}}"] = replacements["{{type_5}}"] = "    "
    elif 3.5 <= avg_point < 4:
        replacements["{{type_2}}"] = "✅"
        replacements["{{type_1}}"] = replacements["{{type_3}}"] = replacements["{{type_4}}"] = replacements["{{type_5}}"] = "    "
    elif 2.5 <= avg_point < 3.5:
        replacements["{{type_3}}"] = "✅"
        replacements["{{type_1}}"] = replacements["{{type_2}}"] = replacements["{{type_4}}"] = replacements["{{type_5}}"] = "    "
    elif 1.5 <= avg_point < 2.5:
        replacements["{{type_4}}"] = "✅"
        replacements["{{type_1}}"] = replacements["{{type_2}}"] = replacements["{{type_3}}"] = replacements["{{type_5}}"] = "    "
    else:
        replacements["{{type_5}}"] = "✅"
        replacements["{{type_1}}"] = replacements["{{type_2}}"] = replacements["{{type_3}}"] = replacements["{{type_4}}"] = "    "


    # Add feedback paragraphs to replacements
    paragraphs = [p.strip() for p in feedback.strip().split("\n\n") if p.strip()]
    print(paragraphs)
    for i in range(1, 9):
        placeholder = f"{{{{feedback_{i+1}}}}}"
        replacements[placeholder] = paragraphs[i - 1] if i <= len(paragraphs) else ""

    print(replacements)

    # Fill text in tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for placeholder, value in replacements.items():
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)
                                run.font.size = Pt(10)

    # Fill text outside tables
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            for placeholder, value in replacements.items():
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)


    output_path = create_output_path(student_name, subject, trial_date, file_extension="docx")
    document.save(output_path)
    print(f"📄 File đã được lưu tại: {output_path}")
    return output_path

